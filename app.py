# --- IMPORTAÇÕES E CONFIGURAÇÃO ---

import streamlit as st
from dotenv import load_dotenv
import os
import re
import pandas as pd
import io
import json
import sys
from fpdf import FPDF
from docx import Document
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain

# Importações opcionais para processamento de arquivos
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Carregar variáveis de ambiente
env_path = ".env"
load_dotenv(env_path)
google_api_key = os.getenv("GOOGLE_API_KEY")

if not google_api_key:
    st.error("Chave de API do Google não encontrada. Certifique-se de definir GOOGLE_API_KEY no .env.")
    st.stop()

llm = ChatGoogleGenerativeAI(
    model="gemini-1.5-flash-latest",
    google_api_key=google_api_key,
    temperature=0.2
)


# --- CHAIN 1: EXTRAÇÃO DE DADOS ---
def get_data_extraction_chain():
    template = """
    Você é um especialista jurídico-financeiro que atua no apoio à administração pública para análise de projetos de lei que envolvem despesas com pessoal, conforme exige a Lei de Responsabilidade Fiscal (LRF).

    Sua tarefa é analisar o texto a seguir e extrair dados estruturados que permitam a elaboração de um estudo de impacto financeiro, considerando aspectos legais, operacionais e financeiros.

    Texto do projeto de lei:
    {text}

    Extraia as seguintes informações no formato JSON. Se algum item não for encontrado, use "Não especificado":

    ```json
    {{
      "tipo_proposta": "",
      "reajuste_proposto": "",
      "abrangencia_temporal": "",
      "setor_afetado": "",
      "detalhes_adicionais": "",
      "quantitativo_envolvido": "",
      "fonte_orcamentaria": "",
      "condicionantes_legais": "",
      "natureza_juridica_da_medida": ""
    }}
    ```
    """
    prompt = PromptTemplate(template=template, input_variables=["text"])
    return LLMChain(llm=llm, prompt=prompt)


# --- CHAIN 2: VALIDAÇÃO LEGAL ---
def get_legal_validation_chain():
    template = """
    Você é um consultor jurídico com foco na Lei de Responsabilidade Fiscal (LRF).

    Analise o seguinte projeto de lei e informe se ele cumpre as exigências legais para aumento de despesa com pessoal:

    Texto:
    {text}

    Responda em formato JSON:
    ```json
    {{
      "cumpre_lrf": "Sim" ou "Não",
      "justificativa": "Explicação concisa sobre o motivo."
    }}
    ```
    """
    prompt = PromptTemplate(template=template, input_variables=["text"])
    return LLMChain(llm=llm, prompt=prompt)


# --- CHAIN 3: AJUSTES SUGERIDOS ---
def get_adjustment_suggestion_chain():
    template = """
    Com base no seguinte texto de projeto de lei, sugira ajustes ou melhorias para garantir conformidade com a LRF e viabilidade financeira:

    Texto:
    {text}

    Responda em formato estruturado:
    ```json
    {{
      "ajustes_sugeridos": [
        "...",
        "..."
      ]
    }}
    ```
    """
    prompt = PromptTemplate(template=template, input_variables=["text"])
    return LLMChain(llm=llm, prompt=prompt)


# --- UTILITÁRIOS ---
def extract_text_from_file(uploaded_file):
    """Extrai texto de diferentes tipos de arquivo"""
    try:
        file_type = uploaded_file.type
        
        # Reset file pointer to beginning
        uploaded_file.seek(0)
        
        if file_type == "text/plain":
            # Arquivo TXT
            try:
                # Try to read as bytes first, then decode
                content = uploaded_file.read()
                if isinstance(content, (bytes, bytearray)):
                    text = content.decode("utf-8")
                else:
                    text = str(content)
                return text
            except UnicodeDecodeError:
                # Try with latin-1 encoding if utf-8 fails
                uploaded_file.seek(0)
                content = uploaded_file.read()
                text = content.decode("latin-1")
                return text
        
        elif file_type == "application/pdf":
            # Arquivo PDF
            if not PDF_AVAILABLE:
                st.error("⚠️ Para ler arquivos PDF, instale: `pip install PyPDF2`")
                return None
                
            try:
                # Create a BytesIO object from uploaded file
                pdf_bytes = uploaded_file.read()
                pdf_file = io.BytesIO(pdf_bytes)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                text = ""
                for page in pdf_reader.pages:
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    except Exception as page_error:
                        st.warning(f"⚠️ Erro ao extrair texto da página: {str(page_error)}")
                        continue
                
                if not text.strip():
                    st.warning("⚠️ Nenhum texto foi extraído do PDF. O arquivo pode conter apenas imagens.")
                    return None
                    
                return text
            except Exception as e:
                st.error(f"❌ Erro ao processar PDF: {str(e)}")
                return None
        
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # Arquivo DOCX
            if not DOCX_AVAILABLE:
                st.error("⚠️ Para ler arquivos DOCX, instale: `pip install python-docx`")
                return None
                
            try:
                # Create a BytesIO object from uploaded file
                docx_bytes = uploaded_file.read()
                docx_file = io.BytesIO(docx_bytes)
                doc = DocxDocument(docx_file)
                
                text = ""
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text += paragraph.text + "\n"
                
                # Also extract text from tables if any
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text += cell.text + "\n"
                
                if not text.strip():
                    st.warning("⚠️ Nenhum texto foi encontrado no documento DOCX.")
                    return None
                    
                return text
            except Exception as e:
                st.error(f"❌ Erro ao processar DOCX: {str(e)}")
                return None
        
        elif file_type == "application/msword":
            # Arquivo DOC (mais limitado)
            st.warning("⚠️ Arquivos .doc não são totalmente suportados. Use .docx ou .txt")
            return None
        
        else:
            st.error(f"⚠️ Tipo de arquivo não suportado: {file_type}")
            st.info("Formatos aceitos: .txt, .pdf, .docx")
            return None
            
    except Exception as e:
        st.error(f"❌ Erro ao processar arquivo: {str(e)}")
        st.error(f"Detalhes técnicos: {type(e).__name__}")
        return None

def extract_percentage(text):
    match = re.search(r'(\d+(\.\d+)?)%', text)
    if match:
        return float(match.group(1))
    return None

def calculate_financial_impact(personnel_expenses, reajuste_percent):
    try:
        reajuste_decimal = reajuste_percent / 100
        impacto_mensal = personnel_expenses * reajuste_decimal
        impacto_anual = impacto_mensal * 12
        return impacto_mensal, impacto_anual
    except:
        return None, None

def parse_llm_response(response_text):
    """Função para extrair JSON da resposta do LLM"""
    try:
        # Remove markdown e espaços
        clean_text = response_text.strip()
        
        # Remove ```json e ``` se existirem
        if "```json" in clean_text:
            start = clean_text.find("```json") + 7
            end = clean_text.rfind("```")
            clean_text = clean_text[start:end]
        elif "```" in clean_text:
            start = clean_text.find("```") + 3
            end = clean_text.rfind("```")
            clean_text = clean_text[start:end]
        
        # Parse JSON
        return json.loads(clean_text.strip())
    except:
        # Se falhar, retorna um dicionário vazio
        return {}

def create_report_text(extracted_data, validacao, sugestoes, reajuste_percent, personnel_expenses, mensal_impact, anual_impact):
    ajustes_str = "\n- ".join(sugestoes.get("ajustes_sugeridos", [])) if sugestoes and sugestoes.get("ajustes_sugeridos") else "Nenhum ajuste sugerido."
    return f"""
Estudo de Impacto Financeiro - Proposta de Reajuste Salarial

1. Descrição da Proposta:
- Tipo: {extracted_data.get('tipo_proposta', 'Não especificado')}
- Setor Afetado: {extracted_data.get('setor_afetado', 'Não especificado')}
- Detalhes Adicionais: {extracted_data.get('detalhes_adicionais', 'Não especificado')}
- Percentual de Reajuste: {reajuste_percent:.2f}%
- Abrangência Temporal: {extracted_data.get('abrangencia_temporal', 'Não especificado')}
- Quantitativo Envolvido: {extracted_data.get('quantitativo_envolvido', 'Não especificado')}
- Fonte Orçamentária: {extracted_data.get('fonte_orcamentaria', 'Não especificado')}
- Condicionantes Legais: {extracted_data.get('condicionantes_legais', 'Não especificado')}
- Natureza Jurídica: {extracted_data.get('natureza_juridica_da_medida', 'Não especificado')}

2. Resultados do Cálculo:
- Gastos Atuais com Pessoal: R$ {personnel_expenses:,.2f}
- Impacto Mensal: R$ {mensal_impact:,.2f}
- Impacto Anual: R$ {anual_impact:,.2f}

3. Validação Jurídica:
- Cumpre LRF? {validacao.get('cumpre_lrf', 'N/A')}
- Justificativa: {validacao.get('justificativa', 'N/A')}

4. Ajustes Sugeridos:
- {ajustes_str}
    """

def create_pdf_report(texto):
    """Cria relatório em PDF com tratamento de encoding robusto"""
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=10)
        
        # Limpar e preparar o texto
        texto_limpo = texto.replace('\r\n', '\n').replace('\r', '\n')
        
        # Tentar diferentes encodings
        try:
            # Primeiro tenta latin-1 (mais compatível com FPDF)
            texto_encoded = texto_limpo.encode('latin-1', 'replace').decode('latin-1')
        except:
            try:
                # Se falhar, tenta utf-8 com replacement
                texto_encoded = texto_limpo.encode('utf-8', 'replace').decode('utf-8')
                # Remove caracteres problemáticos
                texto_encoded = ''.join(char if ord(char) < 256 else '?' for char in texto_encoded)
            except:
                # Último recurso: apenas ASCII
                texto_encoded = ''.join(char if ord(char) < 128 else '?' for char in texto_limpo)
        
        # Adicionar texto ao PDF
        for linha in texto_encoded.split('\n'):
            if linha.strip():
                try:
                    pdf.cell(0, 6, linha[:150], ln=True)  # Limita tamanho da linha
                except:
                    # Se der erro na linha, pula
                    continue
        
        # Gerar PDF
        pdf_output = io.BytesIO()
        pdf_string = pdf.output(dest='S')
        
        if isinstance(pdf_string, str):
            pdf_output.write(pdf_string.encode('latin-1'))
        else:
            pdf_output.write(pdf_string)
            
        pdf_output.seek(0)
        return pdf_output.getvalue()
        
    except Exception as e:
        st.error(f"❌ Erro ao gerar PDF: {str(e)}")
        # Retorna PDF vazio em caso de erro
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(0, 10, "Erro ao gerar relatorio. Use a opcao Word.", ln=True)
        pdf_output = io.BytesIO()
        pdf_string = pdf.output(dest='S')
        if isinstance(pdf_string, str):
            pdf_output.write(pdf_string.encode('latin-1'))
        else:
            pdf_output.write(pdf_string)
        pdf_output.seek(0)
        return pdf_output.getvalue()

def create_word_report(texto):
    doc = Document()
    doc.add_heading('Estudo de Impacto Financeiro', 0)
    
    for linha in texto.strip().split('\n'):
        if linha.strip():
            doc.add_paragraph(linha.strip())
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()

def display_results(dados, validacao, sugestoes, reajuste, gasto_atual, impacto_mensal, impacto_anual):
    """Função para exibir os resultados de forma amigável"""
    
    # Cabeçalho principal
    st.markdown("## 📊 Resultado do Estudo de Impacto Financeiro")
    st.markdown("---")
    
    # 1. Descrição da Proposta
    st.markdown("### 📝 Descrição da Proposta")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.info(f"**Tipo:** {dados.get('tipo_proposta', 'Não especificado')}")
        st.info(f"**Setor Afetado:** {dados.get('setor_afetado', 'Não especificado')}")
        st.info(f"**Percentual de Reajuste:** {reajuste:.2f}%")
        st.info(f"**Abrangência Temporal:** {dados.get('abrangencia_temporal', 'Não especificado')}")
        st.info(f"**Quantitativo Envolvido:** {dados.get('quantitativo_envolvido', 'Não especificado')}")
    
    with col2:
        st.info(f"**Fonte Orçamentária:** {dados.get('fonte_orcamentaria', 'Não especificado')}")
        st.info(f"**Condicionantes Legais:** {dados.get('condicionantes_legais', 'Não especificado')}")
        st.info(f"**Natureza Jurídica:** {dados.get('natureza_juridica_da_medida', 'Não especificado')}")
    
    if dados.get('detalhes_adicionais') and dados.get('detalhes_adicionais') != 'Não especificado':
        st.markdown("**Detalhes Adicionais:**")
        st.write(dados.get('detalhes_adicionais'))
    
    st.markdown("---")
    
    # 2. Impacto Financeiro
    st.markdown("### 💰 Impacto Financeiro")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric(
            label="💼 Gastos Atuais com Pessoal",
            value=f"R$ {gasto_atual:,.2f}"
        )
    
    with col2:
        st.metric(
            label="📅 Impacto Mensal",
            value=f"R$ {impacto_mensal:,.2f}",
            delta=f"{reajuste:.2f}%"
        )
    
    with col3:
        st.metric(
            label="📈 Impacto Anual",
            value=f"R$ {impacto_anual:,.2f}",
            delta=f"R$ {impacto_anual:,.2f}"
        )
    
    # Gráfico de pizza para visualização
    df_impacto = pd.DataFrame({
        'Categoria': ['Gasto Atual', 'Impacto do Reajuste'],
        'Valor': [gasto_atual, impacto_anual]
    })
    
    st.markdown("#### 📊 Visualização do Impacto")
    st.bar_chart(df_impacto.set_index('Categoria'))
    
    st.markdown("---")
    
    # 3. Validação Jurídica
    st.markdown("### ⚖️ Validação Jurídica")
    
    cumpre_lrf = validacao.get('cumpre_lrf', 'N/A')
    
    if cumpre_lrf.lower() == 'sim':
        st.success(f"✅ **Cumpre LRF:** {cumpre_lrf}")
    elif cumpre_lrf.lower() == 'não':
        st.error(f"❌ **Cumpre LRF:** {cumpre_lrf}")
    else:
        st.warning(f"⚠️ **Cumpre LRF:** {cumpre_lrf}")
    
    if validacao.get('justificativa'):
        st.markdown("**Justificativa:**")
        st.write(validacao.get('justificativa'))
    
    st.markdown("---")
    
    # 4. Sugestões de Ajustes
    st.markdown("### 💡 Sugestões de Ajustes")
    
    if sugestoes and sugestoes.get('ajustes_sugeridos'):
        for i, ajuste in enumerate(sugestoes.get('ajustes_sugeridos'), 1):
            st.markdown(f"**{i}.** {ajuste}")
    else:
        st.success("✅ Nenhum ajuste necessário identificado.")
    
    st.markdown("---")


# --- STREAMLIT ---
st.set_page_config(
    page_title="Estudo de Impacto Financeiro", 
    layout="wide",
    page_icon="📊"
)

# Header personalizado
st.markdown("""
<div style='text-align: center; padding: 20px;'>
    <h1 style='color: #2e7d32;'>📊 Sistema de Estudo de Impacto Financeiro</h1>
    <p style='font-size: 18px; color: #666;'>Análise automática de projetos de lei conforme a Lei de Responsabilidade Fiscal (LRF)</p>
</div>
""", unsafe_allow_html=True)

st.markdown("---")

# Sidebar com informações
with st.sidebar:
    st.markdown("## ℹ️ Sobre o Sistema")
    st.markdown("""
    Este sistema realiza:
    - ✅ Extração automática de dados
    - ⚖️ Validação jurídica (LRF)
    - 📊 Cálculo de impacto financeiro
    - 💡 Sugestões de ajustes
    - 📄 Geração de relatórios
    """)
    
    st.markdown("## 📋 Como usar")
    st.markdown("""
    1. **Entrada de dados:**
       - Digite/cole o texto OU
       - Faça upload de arquivo (.txt, .pdf, .docx)
    2. Informe o gasto atual com pessoal
    3. Clique em "Analisar"
    4. Baixe os relatórios se necessário
    """)
    
    st.markdown("## 📁 Formatos Suportados")
    st.markdown("""
    - **TXT:** Texto simples
    - **PDF:** Documentos Adobe
    - **DOCX:** Microsoft Word
    
    ⚠️ **Nota:** Para usar PDF/DOCX, instale:
    ```
    pip install PyPDF2 python-docx
    ```
    """)
    
    st.markdown("---")
    st.markdown("**💡 Dica:** Você pode editar o texto extraído antes da análise!")
# Input principal
st.markdown("## 📝 Entrada de Dados")

# Opções de entrada
input_method = st.radio(
    "Escolha o método de entrada:",
    ["✍️ Digitar texto", "📁 Upload de arquivo"],
    horizontal=True
)

texto = ""

if input_method == "✍️ Digitar texto":
    texto = st.text_area(
        "Cole aqui o texto do projeto de lei:", 
        height=200,
        help="Cole o texto completo do projeto de lei que você deseja analisar"
    )

else:  # Upload de arquivo
    st.markdown("### 📁 Upload de Arquivo")
    
    uploaded_file = st.file_uploader(
        "Escolha um arquivo",
        type=['txt', 'pdf', 'docx'],
        help="Formatos suportados: .txt, .pdf, .docx"
    )
    
    if uploaded_file is not None:
        # Mostrar informações do arquivo
        file_details = {
            "Nome": uploaded_file.name,
            "Tipo": uploaded_file.type,
            "Tamanho": f"{uploaded_file.size / 1024:.2f} KB"
        }
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.info(f"📄 **Nome:** {file_details['Nome']}")
        with col2:
            st.info(f"🔧 **Tipo:** {uploaded_file.type.split('/')[-1].upper()}")
        with col3:
            st.info(f"📊 **Tamanho:** {file_details['Tamanho']}")
        
        # Extrair texto do arquivo
        with st.spinner("📖 Extraindo texto do arquivo..."):
            texto = extract_text_from_file(uploaded_file)
        
        if texto:
            st.success(f"✅ Texto extraído com sucesso! ({len(texto)} caracteres)")
            
            # Mostrar prévia do texto
            with st.expander("👀 Prévia do texto extraído"):
                st.text_area(
                    "Conteúdo do arquivo:",
                    value=texto[:1000] + ("..." if len(texto) > 1000 else ""),
                    height=150,
                    disabled=True
                )
        else:
            st.error("❌ Não foi possível extrair texto do arquivo.")
            
            # Oferecer alternativa manual
            st.info("💡 **Alternativa**: Copie o conteúdo do arquivo e cole na área de texto abaixo:")
            texto_manual = st.text_area(
                "Cole o texto do documento aqui:",
                height=200,
                help="Como alternativa, copie e cole o conteúdo do documento manualmente"
            )
            if texto_manual.strip():
                texto = texto_manual
                st.success("✅ Texto inserido manualmente!")
    
    # Área de texto adicional para edições
    if uploaded_file is not None and texto:
        st.markdown("### ✏️ Edição (Opcional)")
        texto_editado = st.text_area(
            "Você pode editar o texto extraído se necessário:",
            value=texto,
            height=150,
            help="Faça ajustes no texto extraído se necessário"
        )
        if texto_editado != texto:
            texto = texto_editado
            st.info("📝 Texto modificado pelo usuário")

# Validação de entrada
if not texto:
    if input_method == "✍️ Digitar texto":
        st.warning("⚠️ Por favor, digite ou cole o texto do projeto de lei.")
    else:
        st.warning("⚠️ Por favor, faça upload de um arquivo ou digite o texto.")
else:
    # Mostrar estatísticas do texto
    with st.expander("📊 Estatísticas do Texto"):
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Caracteres", len(texto))
        with col2:
            st.metric("Palavras", len(texto.split()))
        with col3:
            st.metric("Parágrafos", len([p for p in texto.split('\n') if p.strip()]))
        with col4:
            # Estimar porcentagem se encontrada
            perc = extract_percentage(texto)
            st.metric("% Encontrado", f"{perc}%" if perc else "N/A")

# Configurações adicionais
with st.expander("⚙️ Configurações Avançadas"):
    gasto_atual = st.number_input(
        "Gasto Atual com Pessoal (R$):", 
        value=10000000.0, 
        step=10000.0,
        help="Informe o valor atual gasto com pessoal para cálculo do impacto"
    )

# Botão de análise
if st.button("🔍 Analisar e Gerar Estudo", type="primary", disabled=not texto):
    if texto:
        with st.spinner("🔄 Executando análise completa..."):
            try:
                # Executar as chains
                progress_bar = st.progress(0)
                
                # Chain 1 - Extração de dados
                st.write("📝 Extraindo dados do projeto...")
                progress_bar.progress(25)
                dados_response = get_data_extraction_chain().run(text=texto)
                dados = parse_llm_response(dados_response)
                
                # Chain 2 - Validação legal
                st.write("⚖️ Realizando validação jurídica...")
                progress_bar.progress(50)
                validacao_response = get_legal_validation_chain().run(text=texto)
                validacao = parse_llm_response(validacao_response)
                
                # Chain 3 - Sugestões
                st.write("💡 Gerando sugestões de ajustes...")
                progress_bar.progress(75)
                sugestoes_response = get_adjustment_suggestion_chain().run(text=texto)
                sugestoes = parse_llm_response(sugestoes_response)
                
                # Cálculos
                st.write("📊 Calculando impacto financeiro...")
                progress_bar.progress(100)
                
                reajuste = extract_percentage(dados.get("reajuste_proposto", "0%")) or 5.0
                impacto_mensal, impacto_anual = calculate_financial_impact(gasto_atual, reajuste)
                
                # Exibir resultados
                display_results(dados, validacao, sugestoes, reajuste, gasto_atual, impacto_mensal, impacto_anual)
                
                # Botões de download
                st.markdown("## 📥 Downloads")
                
                texto_relatorio = create_report_text(
                    dados, validacao, sugestoes, reajuste, 
                    gasto_atual, impacto_mensal, impacto_anual
                )
                
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.download_button(
                        label="📄 Baixar PDF",
                        data=create_pdf_report(texto_relatorio),
                        file_name="relatorio_impacto.pdf",
                        mime="application/pdf"
                    )
                
                with col2:
                    st.download_button(
                        label="📝 Baixar Word",
                        data=create_word_report(texto_relatorio),
                        file_name="relatorio_impacto.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                
                with col3:
                    st.download_button(
                        label="📊 Baixar Dados (JSON)",
                        data=json.dumps({
                            "dados": dados,
                            "validacao": validacao,
                            "sugestoes": sugestoes,
                            "impacto_financeiro": {
                                "reajuste_percent": reajuste,
                                "gasto_atual": gasto_atual,
                                "impacto_mensal": impacto_mensal,
                                "impacto_anual": impacto_anual
                            }
                        }, indent=2, ensure_ascii=False),
                        file_name="dados_analise.json",
                        mime="application/json"
                    )
                
            except Exception as e:
                st.error(f"❌ Erro durante a análise: {str(e)}")
                st.error("Verifique se a API do Google está funcionando corretamente.")
                
                # Debug info
                with st.expander("🔧 Informações de Debug"):
                    st.write(f"Tipo de erro: {type(e).__name__}")
                    st.write(f"Tamanho do texto: {len(texto)} caracteres")
                    st.write(f"Método de entrada: {input_method}")
    else:
        st.error("⚠️ Nenhum texto foi fornecido para análise.")

# Footer
st.markdown("---")
st.markdown("""
<div style='text-align: center; color: #666; padding: 20px;'>
    <small>Sistema de Análise de Impacto Financeiro | Desenvolvido para auxiliar na análise de projetos de lei conforme a LRF</small>
</div>
""", unsafe_allow_html=True)

# Debug info for Streamlit Share
if st.checkbox("🔧 Mostrar informações de debug", value=False):
    st.markdown("### 🔧 Informações de Debug")
    st.write(f"**Streamlit version**: {st.__version__}")
    st.write(f"**Python version**: {sys.version}")
    
    # Test dependencies
    deps_status = {}
    try:
        import PyPDF2
        deps_status['PyPDF2'] = f"✅ {PyPDF2.__version__}"
    except:
        deps_status['PyPDF2'] = "❌ Não disponível"
    
    try:
        from docx import __version__ as docx_version
        deps_status['python-docx'] = f"✅ {docx_version}"
    except:
        try:
            import docx
            deps_status['python-docx'] = "✅ Instalado (versão desconhecida)"
        except:
            deps_status['python-docx'] = "❌ Não disponível"
    
    try:
        from fpdf import __version__ as fpdf_version
        deps_status['fpdf2'] = f"✅ {fpdf_version}"
    except:
        try:
            import fpdf
            deps_status['fpdf2'] = "✅ Instalado (versão desconhecida)"
        except:
            deps_status['fpdf2'] = "❌ Não disponível"
    
    for dep, status in deps_status.items():
        st.write(f"**{dep}**: {status}")